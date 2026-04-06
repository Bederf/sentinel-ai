"""Document text extraction service for multiple file types."""

import io
import logging
from pathlib import Path

from fastapi import UploadFile

logger = logging.getLogger(__name__)

PDF_LOW_TEXT_THRESHOLD = 200


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2.

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        Extracted text content
    """
    try:
        from io import BytesIO

        from PyPDF2 import PdfReader

        pdf_reader = PdfReader(BytesIO(file_bytes))
        text_parts = []

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")

        return "\n\n".join(text_parts)

    except ImportError as exc:
        raise ImportError("PyPDF2 is required for PDF extraction. Install with: pip install PyPDF2") from exc
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise


def _extract_text_from_pdf_native(file_bytes: bytes) -> str:
    """Native PyPDF2 extraction — no OCR involved."""
    return extract_text_from_pdf(file_bytes)


def extract_text_from_pdf_docling(file_bytes: bytes) -> str:
    """Extract text from PDF using Docling OCR/layout analysis.

    Docling provides superior scanned-PDF OCR, table extraction, and key-value
    pair detection compared to PyMuPDF+pytesseract.

    Returns an empty string if docling is unavailable or extraction fails.
    """
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        logger.debug("Docling unavailable for PDF OCR")
        return ""

    try:
        from io import BytesIO

        converter = DocumentConverter()
        result = converter.convert(BytesIO(file_bytes))
        text = result.document.export_to_markdown()
        return text.strip() if text else ""
    except Exception as exc:
        logger.info("Docling PDF extraction failed: %s", exc)
        return ""


def _extract_text_from_pdf_ocr(file_bytes: bytes) -> str:
    """OCR fallback for low-text or scanned PDFs using PyMuPDF+pytesseract.

    Returns an empty string if OCR dependencies are unavailable or OCR fails.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except Exception as exc:
        logger.debug("PDF OCR fallback unavailable: %s", exc)
        return ""

    text_parts: list[str] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            text_parts.append(pytesseract.image_to_string(image, lang="eng"))
        return "\n\n".join(part.strip() for part in text_parts if part and part.strip())
    except Exception as exc:
        logger.info("PDF OCR fallback failed: %s", exc)
        return ""


def extract_text_from_pdf_with_fallback(file_bytes: bytes) -> tuple[str, dict]:
    """Extract PDF text: PyPDF2 native -> Docling OCR (if <200 chars) -> pytesseract."""
    native_text = _extract_text_from_pdf_native(file_bytes)
    native_length = len(native_text.strip())

    metadata: dict = {
        "file_type": ".pdf",
        "native_text_length": native_length,
        "low_text_threshold": PDF_LOW_TEXT_THRESHOLD,
        "ocr_used": False,
    }

    if native_length >= PDF_LOW_TEXT_THRESHOLD:
        metadata["extraction_mode"] = "native"
        return native_text, metadata

    # Try docling OCR first (better than pytesseract for scanned/image PDFs)
    docling_text = extract_text_from_pdf_docling(file_bytes)
    if docling_text.strip():
        metadata["ocr_used"] = True
        metadata["extraction_mode"] = "docling"
        metadata["ocr_text_length"] = len(docling_text.strip())
        metadata["fallback_reason"] = "low_native_text_pdf"
        return docling_text, metadata

    # Final fallback: pytesseract
    ocr_text = _extract_text_from_pdf_ocr(file_bytes)
    if ocr_text.strip():
        metadata["ocr_used"] = True
        metadata["extraction_mode"] = "ocr_fallback"
        metadata["ocr_text_length"] = len(ocr_text.strip())
        metadata["fallback_reason"] = "low_native_text_pdf"
        return ocr_text, metadata

    metadata["extraction_mode"] = "native_low_text_no_ocr"
    metadata["fallback_reason"] = "ocr_unavailable_or_failed"
    return native_text, metadata


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx.

    Args:
        file_bytes: Raw DOCX file bytes

    Returns:
        Extracted text content
    """
    try:
        from io import BytesIO

        from docx import Document

        doc = Document(BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n\n".join(text_parts)

    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx") from exc
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise


async def extract_text(file: UploadFile) -> tuple[str, dict]:
    """Extract text from uploaded file based on file type.

    Supports: PDF, DOCX, TXT

    Args:
        file: UploadFile from FastAPI multipart form

    Returns:
        Tuple of (extracted_text, metadata_dict)

    Raises:
        ValueError: If file type is not supported
        ImportError: If required package is not installed
    """
    # Read file content
    content = await file.read()

    # Determine file type from extension
    file_ext = Path(file.filename or "").suffix.lower()

    if file_ext == ".pdf":
        text, pdf_meta = extract_text_from_pdf_with_fallback(content)
    elif file_ext == ".docx":
        text = extract_text_from_docx(content)
    elif file_ext == ".txt":
        text = content.decode("utf-8")
    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Supported types: .pdf, .docx, .txt")

    # Build metadata
    metadata = {
        "filename": file.filename,
        "size_bytes": len(content),
        "file_type": file_ext,
    }
    if file_ext == ".pdf":
        metadata.update(pdf_meta)

    logger.info(f"Extracted text from {file.filename}: {len(text)} characters")

    return text, metadata
